from pydoc import doc
import docx

# create a new Word document
document = docx.Document()

# add a heading
document.add_heading('Sean McNamara', 0)

# Add a paragraph under the heading
document.add_paragraph('“Programming languages such as JavaScript and Python, web development using technologies such as HTML, CSS, and JavaScript frameworks like VueJS”')

# add personal information
personal_info = document.add_paragraph('Personal Information:')
personal_info.add_run('\nName: ').bold = True
personal_info.add_run('Sean McNamara')
personal_info.add_run('\nEmail: ').bold = True
personal_info.add_run('seanmcnamara812@gmai.com')
personal_info.add_run('\nPhone: ').bold = True
personal_info.add_run('07572603965')

# add education information
education = document.add_paragraph('Education:')
education.add_run('\nCode Lancashire').bold = True
education.add_run('\nBoot camp 2023')

# add work experience
experience = document.add_paragraph('Work Experience:')
company = experience.add_run('\nRetraining to be in software development or data science')
company.bold = True
company.italic = True
experience.add_run('\nJan - Mar 2023')
experience.add_run('\n- Developed Pokedex App using VueJs and MySql')
experience.add_run('\n- Learnt how to use Python for data analysis using Numpy and Pandas')
experience.add_run('\n-Enhancing my soft skills including problem-solving and critical thinking')

# save the document
document.save('cv.docx')
